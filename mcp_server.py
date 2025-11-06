#!/usr/bin/env python3
"""
Complete MCP Server for Content & Media Processing
(*** ROBUST LOCAL PANDOC VERSION - 11/06 - HANG FIX ***)
"""

import asyncio
import json
import sys
import os
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Sequence
from datetime import datetime
import tempfile
import shutil
import uuid
import sqlite3

# Core MCP imports
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp import types

# --- Processing libraries ---

# For Image Processing
from PIL import Image, ImageEnhance

# For News Processing
import feedparser
import requests 

# For Email Processing
import email
import mailbox
from email.parser import BytesParser
from email.policy import default

# For File Type Detection
import re
import magic

# --- Local Document Conversion ---
import pypandoc
import PyPDF2
import docx  # python-docx
import pptx  # python-pptx

# --- API and Environment Imports ---
# import cloudconvert # No longer needed
from dotenv import load_dotenv
import google.generativeai as genai 

# --- NEW: Configure pypandoc to NOT download binaries ---
# This forces it to use the system-installed pandoc and fail fast if it's missing.
os.environ['PYPANDOC_DONT_DOWNLOAD'] = '1' 

# Load .env file for API keys
load_dotenv()

# Configuration
class Config:
    BASE_DIR = Path(__file__).parent
    TEMP_DIR = BASE_DIR / "temp"
    OUTPUT_DIR = BASE_DIR / "output" 
    CACHE_DIR = BASE_DIR / "cache"
    DB_PATH = BASE_DIR / "mcp_server.db"
    
    # File size limits (in bytes)
    MAX_DOC_SIZE = 50 * 1024 * 1024  # 50MB
    MAX_IMG_SIZE = 100 * 1024 * 1024  # 100MB
    MAX_EMAIL_SIZE = 5 * 1024 * 1024 * 1024  # 5GB
    
    # Supported formats
    DOC_FORMATS = ["pdf", "docx", "doc", "html", "htm", "md", "markdown", "txt", "rtf", "pptx"]
    IMG_FORMATS = ["jpg", "jpeg", "png", "gif", "bmp", "tiff", "webp"]
    EMAIL_FORMATS = ["mbox", "eml"]
    
    def __post_init__(self):
        for directory in [self.TEMP_DIR, self.OUTPUT_DIR, self.CACHE_DIR]:
            directory.mkdir(parents=True, exist_ok=True)

config = Config()
config.__post_init__()

# Logging setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize MCP Server
app = Server("content-media-hub")

# --- Database Setup (CORRECTED) ---
def init_database():
    """Initialize database tables"""
    try:
        with get_db_connection() as conn:
            # Processing history table
            conn.execute("""
                CREATE TABLE IF NOT EXISTS processing_history (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    tool_name TEXT NOT NULL,
                    source_path TEXT,
                    output_path TEXT,
                    status TEXT NOT NULL,
                    error_message TEXT,
                    metadata TEXT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # News sources table
            conn.execute("""
                CREATE TABLE IF NOT EXISTS news_sources (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    url TEXT UNIQUE NOT NULL,
                    title TEXT,
                    description TEXT,
                    category TEXT,
                    status TEXT DEFAULT 'active',
                    last_updated DATETIME,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # News articles table
            conn.execute("""
                CREATE TABLE IF NOT EXISTS news_articles (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    source_id INTEGER,
                    title TEXT,
                    summary TEXT,
                    url TEXT,
                    published_date TEXT,
                    hash TEXT UNIQUE,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (source_id) REFERENCES news_sources (id)
                )
            """)
            
            # --- START: BUG FIX ---
            # Added UNIQUE constraint to fix the "ON CONFLICT" error
            conn.execute("""
                CREATE TABLE IF NOT EXISTS trending_keywords (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    keyword TEXT,
                    count INTEGER,
                    category TEXT,
                    time_window TEXT,
                    last_seen DATETIME DEFAULT CURRENT_TIMESTAMP,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE(keyword, category, time_window) 
                )
            """)
            # --- END: BUG FIX ---
            
            conn.commit()
        logger.info("Database initialized successfully")
    except Exception as e:
        logger.error(f"Database initialization failed: {e}")

# --- Database Utility Functions ---
def get_db_connection():
    """Create a database connection"""
    conn = sqlite3.connect(config.DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def log_to_history(tool_name: str, status: str, source_path: str = None, output_path: str = None, error_message: str = None, metadata: dict = None):
    """Log an action to the processing_history table"""
    try:
        with get_db_connection() as conn:
            conn.execute(
                """
                INSERT INTO processing_history (tool_name, source_path, output_path, status, error_message, metadata)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (tool_name, source_path, output_path, status, error_message, json.dumps(metadata) if metadata else None)
            )
            conn.commit() # Make sure to commit the changes
    except Exception as e:
        logger.error(f"Failed to log to database: {e}")

# --- Utility Functions ---
def ensure_directory(path: Path):
    """Ensure directory exists"""
    path.mkdir(parents=True, exist_ok=True)

def get_file_hash(filepath: str) -> str:
    """Generate hash for file"""
    import hashlib
    return hashlib.md5(filepath.encode()).hexdigest()[:8]

def validate_file_size(filepath: Path, max_size: int):
    """Validate file size"""
    if filepath.stat().st_size > max_size:
        raise ValueError(f"File too large: {filepath.stat().st_size / (1024*1024):.1f}MB")

def detect_file_format(filepath: Path) -> str:
    """Detect file format"""
    ext = filepath.suffix.lower().lstrip('.')
    try:
        mime_type = magic.from_file(str(filepath), mime=True)
        if 'pdf' in mime_type:
            return 'pdf'
        elif 'word' in mime_type or 'officedocument.wordprocessingml' in mime_type:
            return 'docx' if 'openxml' in mime_type else 'doc'
        elif 'powerpoint' in mime_type or 'officedocument.presentationml' in mime_type:
            return 'pptx'
        elif 'html' in mime_type:
            return 'html'
        elif 'rtf' in mime_type:
            return 'rtf'
        elif 'markdown' in mime_type or ext == 'md':
            return 'markdown'
        elif 'image' in mime_type:
            return ext if ext in config.IMG_FORMATS else 'jpg'
    except Exception as e:
        logger.warning(f"File magic check failed: {e}. Falling back to extension.")
        pass
    
    # Fallback to extension if magic fails
    if ext in config.DOC_FORMATS or ext in config.IMG_FORMATS:
        return ext
    else:
        return 'txt'


# --- Document Processing Service (LOCAL PANDOC VERSION - HANG FIX) ---
class DocumentProcessor:

    """

    Robust document processor using Pypandoc for local file conversion.

    Requires Pandoc system package to be installed on the machine.

    

    Installation:

        - Windows: choco install pandoc

        - macOS: brew install pandoc

        - Linux: sudo apt-get install pandoc

    """

    

    # Supported input formats and their MIME types

    FORMAT_MIME_TYPES = {

        'pdf': 'application/pdf',

        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',

        'doc': 'application/msword',

        'html': 'text/html',

        'htm': 'text/html',

        'md': 'text/markdown',

        'markdown': 'text/markdown',

        'txt': 'text/plain',

        'rtf': 'application/rtf',

        'odt': 'application/vnd.oasis.opendocument.text',

        'tex': 'application/x-latex',

        'latex': 'application/x-latex',

        'epub': 'application/epub+zip',

    }

    

    # Pandoc format names (some differ from file extensions)

    PANDOC_FORMAT_MAP = {

        'docx': 'docx',

        'doc': 'docx',  # Convert .doc to docx first

        'pdf': 'pdf',

        'html': 'html',

        'htm': 'html',

        'md': 'markdown',

        'markdown': 'markdown',

        'txt': 'plain',

        'rtf': 'rtf',

        'odt': 'odt',

        'tex': 'latex',

        'latex': 'latex',

        'epub': 'epub',

    }

    

    # Format compatibility matrix (source -> acceptable targets)

    FORMAT_COMPATIBILITY = {

        'pdf': ['docx', 'txt', 'html', 'md'],

        'docx': ['pdf', 'txt', 'html', 'md', 'odt', 'rtf', 'epub'],

        'doc': ['docx', 'pdf', 'txt', 'html', 'odt'],

        'html': ['pdf', 'docx', 'txt', 'md', 'odt', 'epub'],

        'txt': ['pdf', 'docx', 'html', 'md', 'odt'],

        'md': ['pdf', 'docx', 'html', 'txt', 'odt', 'epub', 'rtf'],

        'rtf': ['pdf', 'docx', 'txt', 'html', 'odt'],

        'odt': ['pdf', 'docx', 'html', 'txt', 'md'],

        'tex': ['pdf', 'html', 'txt', 'docx'],

        'epub': ['pdf', 'docx', 'html', 'txt', 'md'],

    }

    

    # Retry settings

    MAX_RETRIES = 2

    RETRY_DELAY = 1  # seconds

    

    def __init__(self):

        """Initialize DocumentProcessor and verify Pandoc installation"""

        self.output_dir = config.OUTPUT_DIR/ "documents"

        self.temp_dir = config.TEMP_DIR

        self.pandoc_available = False

        

        # Verify Pandoc is installed

        self._verify_pandoc_installation()

        logger.info("✓ DocumentProcessor initialized with Pypandoc")

    

    def _verify_pandoc_installation(self) -> None:

        """Verify that Pandoc is installed and accessible"""

        try:

            if pypandoc is None:

                raise ImportError("pypandoc is not installed")

            

            version = pypandoc.get_pandoc_version()

            self.pandoc_available = True

            logger.info(f"✓ Pandoc version detected: {version}")

        except Exception as e:

            logger.error(f"✗ Pandoc verification failed: {e}")

            raise RuntimeError(

                "Pandoc is required but not installed or not accessible.\n"

                "Please install Pandoc:\n"

                "  - Windows: choco install pandoc\n"

                "  - macOS: brew install pandoc\n"

                "  - Linux: sudo apt-get install pandoc\n"

                f"Error details: {e}"

            )

    

    async def convert_document(self, source_path: str, target_format: str, **kwargs) -> Dict[str, Any]:

        """

        Convert document from one format to another using Pandoc.

        

        Args:

            source_path: Path to source document

            target_format: Target format (e.g., 'pdf', 'docx', 'html', 'md')

            **kwargs: Additional options like:

                - toc: bool (add table of contents)

                - number_sections: bool (number sections)

                - standalone: bool (produce standalone file)

                - pdf_engine: str ('pdflatex', 'xelatex', 'wkhtmltopdf')

                - self_contained: bool (embed resources in HTML)

                - preserve_tabs: bool (preserve tabs in conversion)

        

        Returns:

            Dict containing:

                - success: bool

                - message: str (description)

                - output_path: str (path to converted file)

                - metadata: dict (file info)

                - source_format: str

                - output_format: str

                - file_size: int (bytes)

                - error: str (if failed)

                - error_type: str (type of error)

        """

        try:

            # Validate inputs

            source = self._validate_source_file(source_path)

            source_format = self._detect_format(source)

            target_format = target_format.lower().strip('.')

            

            logger.info(f"Starting conversion: {source.name} ({source_format}) → {target_format}")

            

            # Validate format compatibility

            self._validate_format_compatibility(source_format, target_format)

            

            # Generate output path

            output_path = self._generate_output_path(source, target_format)

            

            # Perform conversion with retry logic

            result = await self._retry_conversion(

                source, output_path, source_format, target_format, **kwargs

            )

            

            if not result.get('success'):

                return result

            

            # Verify output file

            if not output_path.exists():

                raise RuntimeError("Conversion completed but output file not found")

            

            file_size = output_path.stat().st_size

            if file_size == 0:

                raise RuntimeError("Conversion produced an empty file")

            

            # Extract metadata

            metadata = self._extract_metadata(output_path, target_format)

            

            result.update({

                'metadata': metadata,

                'source_format': source_format,

                'output_format': target_format,

                'file_size': file_size

            })

            

            logger.info(f"✓ Successfully converted to {output_path.name} ({file_size} bytes)")

            return result

            

        except FileNotFoundError as e:

            logger.error(f"File not found: {e}")

            return {

                'success': False,

                'error': f"Source file not found: {source_path}",

                'error_type': 'file_not_found',

                'source_path': source_path

            }

        except ValueError as e:

            logger.error(f"Validation error: {e}")

            return {

                'success': False,

                'error': str(e),

                'error_type': 'validation_error',

                'source_path': source_path,

                'target_format': target_format

            }

        except RuntimeError as e:

            logger.error(f"Conversion error: {e}")

            return {

                'success': False,

                'error': str(e),

                'error_type': 'pandoc_error',

                'source_path': source_path,

                'target_format': target_format

            }

        except Exception as e:

            logger.exception(f"Unexpected error during conversion")

            return {

                'success': False,

                'error': str(e),

                'error_type': 'conversion_error',

                'source_path': source_path,

                'target_format': target_format

            }

    

    async def _retry_conversion(self, source: Path, output_path: Path, source_format: str,

                               target_format: str, **kwargs) -> Dict[str, Any]:

        """Execute conversion with retry logic"""

        last_error = None

        

        for attempt in range(self.MAX_RETRIES):

            try:

                loop = asyncio.get_event_loop()

                result = await loop.run_in_executor(

                    None,

                    self._execute_pypandoc_conversion,

                    source,

                    output_path,

                    source_format,

                    target_format,

                    kwargs

                )

                return result

            except Exception as e:

                last_error = e

                logger.warning(f"Conversion attempt {attempt + 1}/{self.MAX_RETRIES} failed: {e}")

                if attempt < self.MAX_RETRIES - 1:

                    await asyncio.sleep(self.RETRY_DELAY * (attempt + 1))

        

        # All retries failed

        raise last_error

    

    def _execute_pypandoc_conversion(self, source: Path, output_path: Path,

                                    source_format: str, target_format: str,

                                    kwargs: Dict) -> Dict[str, Any]:

        """

        Execute Pypandoc conversion synchronously.

        Runs in executor to prevent blocking async code.

        """

        if not self.pandoc_available:

            raise RuntimeError("Pandoc is not available")

        

        try:

            # Get Pandoc format names

            input_format = self.PANDOC_FORMAT_MAP.get(source_format, source_format)

            output_format = self.PANDOC_FORMAT_MAP.get(target_format, target_format)

            

            # Build extra arguments for Pandoc

            extra_args = self._build_pandoc_args(target_format, kwargs)

            

            logger.debug(

                f"Pandoc execution: {input_format} → {output_format}\n"

                f"  Source: {source}\n"

                f"  Output: {output_path}\n"

                f"  Args: {extra_args}"

            )

            

            # Read source file as binary

            with open(source, 'rb') as f:

                source_content = f.read()

            

            # Convert using Pypandoc

            output_content = pypandoc.convert_text(

                source=source_content,

                format=input_format,

                to=output_format,

                extra_args=extra_args,

                outputfile=str(output_path)

            )

            

            # Verify output file was created

            if not output_path.exists():

                raise RuntimeError(

                    f"Pandoc failed to create output file at {output_path}\n"

                    f"Pandoc output: {output_content}"

                )

            

            return {

                'success': True,

                'message': f"Successfully converted {source.name} to {target_format}",

                'output_path': str(output_path),

                'converter': 'pypandoc (Pandoc)'

            }

            

        except Exception as e:

            logger.error(f"Pypandoc conversion error: {e}")

            raise RuntimeError(f"Pandoc conversion failed: {str(e)}")

    

    def _build_pandoc_args(self, target_format: str, kwargs: Dict) -> List[str]:

        """

        Build extra arguments for Pandoc command line.

        

        Reference: https://pandoc.org/MANUAL.html

        """

        args = []

        

        # Standalone option (useful for HTML, LaTeX, Markdown)

        if kwargs.get('standalone', True):

            args.append('--standalone')

        

        # Table of contents

        if kwargs.get('toc', False):

            args.append('--toc')

            # TOC depth

            toc_depth = kwargs.get('toc_depth', 3)

            args.extend(['--toc-depth', str(toc_depth)])

        

        # Number sections

        if kwargs.get('number_sections', False):

            args.append('--number-sections')

        

        # Preserve tabs (don't convert to spaces)

        if kwargs.get('preserve_tabs', False):

            args.append('--preserve-tabs')

        

        # PDF engine selection for PDF output

        if target_format == 'pdf':

            pdf_engine = kwargs.get('pdf_engine', 'pdflatex')

            args.extend(['--pdf-engine', pdf_engine])

            

            # PDF engine options

            if kwargs.get('pdf_engine_opt'):

                args.extend(['--pdf-engine-opt', kwargs['pdf_engine_opt']])

        

        # Reference document (for DOCX output styling)

        if kwargs.get('reference_doc'):

            ref_path = Path(kwargs['reference_doc'])

            if ref_path.exists():

                args.extend(['--reference-doc', str(ref_path)])

            else:

                logger.warning(f"Reference document not found: {ref_path}")

        

        # Template file

        if kwargs.get('template'):

            template_path = Path(kwargs['template'])

            if template_path.exists():

                args.extend(['--template', str(template_path)])

            else:

                logger.warning(f"Template not found: {template_path}")

        

        # Metadata (key=value pairs)

        if kwargs.get('metadata') and isinstance(kwargs['metadata'], dict):

            for key, value in kwargs['metadata'].items():

                args.extend(['-M', f'{key}={value}'])

        

        # Self-contained HTML (embed images, CSS, etc.)

        if kwargs.get('self_contained', False) and target_format == 'html':

            args.append('--self-contained')

        

        # CSS file for HTML output

        if kwargs.get('css') and target_format == 'html':

            css_path = Path(kwargs['css'])

            if css_path.exists():

                args.extend(['-c', str(css_path)])

            else:

                logger.warning(f"CSS file not found: {css_path}")

        

        # Code highlighting

        if kwargs.get('highlight_style'):

            args.extend(['--highlight-style', kwargs['highlight_style']])

        

        # Data directory (for templates, etc.)

        if kwargs.get('data_dir'):

            data_dir = Path(kwargs['data_dir'])

            if data_dir.exists():

                args.extend(['--data-dir', str(data_dir)])

        

        # Verbose output

        if kwargs.get('verbose', False):

            args.append('--verbose')

        

        # Strip empty paragraphs

        if kwargs.get('strip_empty_paragraphs', False):

            args.append('--strip-empty-paragraphs')

        

        # Slide format options

        if target_format in ['pptx', 'beamer']:

            if kwargs.get('slide_level'):

                args.extend(['--slide-level', str(kwargs['slide_level'])])

        

        logger.debug(f"Pandoc arguments: {args}")

        return args

    

    def _validate_source_file(self, source_path: str) -> Path:

        """Validate source file exists and is readable"""

        source = Path(source_path)

        

        # Try direct path first

        if not source.exists():

            # Try relative to base directory

            source = config.BASE_DIR / source.name

            if not source.exists():

                raise FileNotFoundError(f"Source file not found: {source_path}")

        

        # Resolve to absolute path

        source = source.resolve()

        

        # Check file size

        try:

            file_size = source.stat().st_size

        except OSError as e:

            raise PermissionError(f"Cannot access file: {e}")

        

        if file_size == 0:

            raise ValueError("Source file is empty")

        

        if file_size > config.MAX_DOC_SIZE:

            size_mb = file_size / (1024 * 1024)

            max_mb = config.MAX_DOC_SIZE / (1024 * 1024)

            raise ValueError(f"File too large: {size_mb:.1f}MB (max: {max_mb:.0f}MB)")

        

        # Check readability

        if not os.access(source, os.R_OK):

            raise PermissionError(f"File is not readable: {source}")

        

        logger.debug(f"Source file validated: {source} ({file_size} bytes)")

        return source

    

    def _detect_format(self, file_path: Path) -> str:

        """Detect file format from extension and magic bytes"""

        ext = file_path.suffix.lower().lstrip('.')

        

        # Check extension first

        if ext in self.FORMAT_MIME_TYPES:

            logger.debug(f"Format detected by extension: {ext}")

            return ext

        

        # Try magic bytes detection

        try:

            import magic

            mime_type = magic.from_file(str(file_path), mime=True)

            logger.debug(f"MIME type detected: {mime_type}")

            

            # Map MIME type to format

            mime_to_format = {

                'application/pdf': 'pdf',

                'application/msword': 'doc',

                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',

                'text/html': 'html',

                'text/plain': 'txt',

                'text/markdown': 'md',

                'application/vnd.oasis.opendocument.text': 'odt',

                'application/epub+zip': 'epub',

                'application/rtf': 'rtf',

            }

            

            for mime, fmt in mime_to_format.items():

                if mime in mime_type:

                    logger.debug(f"Format detected by MIME type: {fmt}")

                    return fmt

        except Exception as e:

            logger.debug(f"Magic detection failed: {e}")

        

        # Fallback to extension

        if ext and ext in config.DOC_FORMATS:

            logger.debug(f"Format detected by extension fallback: {ext}")

            return ext

        

        raise ValueError(f"Unable to detect file format for: {file_path}")

    

    def _validate_format_compatibility(self, source_format: str, target_format: str) -> None:

        """Validate that conversion is supported"""

        if source_format not in self.FORMAT_COMPATIBILITY:

            raise ValueError(f"Unsupported source format: {source_format}")

        

        compatible_targets = self.FORMAT_COMPATIBILITY.get(source_format, [])

        if target_format not in compatible_targets:

            raise ValueError(

                f"Cannot convert {source_format} to {target_format}. "

                f"Supported targets for {source_format}: {', '.join(compatible_targets)}"

            )

    

    def _generate_output_path(self, source: Path, target_format: str) -> Path:

        """Generate output file path with timestamp"""

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        output_name = f"{source.stem}_{timestamp}.{target_format}"

        return self.output_dir / output_name

    

    def _extract_metadata(self, file_path: Path, file_format: str) -> Dict[str, Any]:

        """Extract file metadata"""

        try:

            stats = file_path.stat()

        except OSError as e:

            logger.warning(f"Cannot read file stats: {e}")

            return {}

        

        metadata = {

            'file_size': stats.st_size,

            'created_time': datetime.fromtimestamp(stats.st_ctime).isoformat(),

            'modified_time': datetime.fromtimestamp(stats.st_mtime).isoformat(),

            'format': file_format,

            'mime_type': self.FORMAT_MIME_TYPES.get(file_format, 'application/octet-stream'),

            'file_name': file_path.name

        }

        

        # Extract format-specific metadata

        try:

            if file_format == 'pdf':

                metadata.update(self._extract_pdf_metadata(file_path))

            elif file_format in ['docx', 'doc']:

                metadata.update(self._extract_docx_metadata(file_path))

            elif file_format == 'html':

                metadata.update(self._extract_html_metadata(file_path))

        except Exception as e:

            logger.debug(f"Format-specific metadata extraction failed: {e}")

        

        return metadata

    

    def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:

        """Extract PDF-specific metadata"""

        try:

            import PyPDF2

            with open(file_path, 'rb') as f:

                reader = PyPDF2.PdfReader(f)

                return {

                    'pages': len(reader.pages),

                    'encrypted': reader.is_encrypted

                }

        except ImportError:

            logger.debug("PyPDF2 not installed, skipping PDF metadata")

            return {}

        except Exception as e:

            logger.debug(f"PDF metadata extraction failed: {e}")

            return {}

    

    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:

        """Extract DOCX-specific metadata"""

        try:

            from docx import Document

            doc = Document(file_path)

            image_count = len([rel for rel in doc.part.rels.values() if 'image' in rel.reltype])

            return {

                'paragraphs': len(doc.paragraphs),

                'tables': len(doc.tables),

                'images': image_count

            }

        except ImportError:

            logger.debug("python-docx not installed, skipping DOCX metadata")

            return {}

        except Exception as e:

            logger.debug(f"DOCX metadata extraction failed: {e}")

            return {}

    

    def _extract_html_metadata(self, file_path: Path) -> Dict[str, Any]:

        """Extract HTML-specific metadata"""

        try:

            from html.parser import HTMLParser

            

            class MetadataParser(HTMLParser):

                def __init__(self):

                    super().__init__()

                    self.title = None

                    self.headings = 0

                    self.paragraphs = 0

                    self.links = 0

                    self.images = 0

                    self.in_title = False

                

                def handle_starttag(self, tag, attrs):

                    if tag == 'title':

                        self.in_title = True

                    elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:

                        self.headings += 1

                    elif tag == 'p':

                        self.paragraphs += 1

                    elif tag == 'a':

                        self.links += 1

                    elif tag == 'img':

                        self.images += 1

                

                def handle_endtag(self, tag):

                    if tag == 'title':

                        self.in_title = False

                

                def handle_data(self, data):

                    if self.in_title and self.title is None:

                        self.title = data.strip()

            

            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:

                parser = MetadataParser()

                parser.feed(f.read())

            

            return {

                'title': parser.title,

                'headings': parser.headings,

                'paragraphs': parser.paragraphs,

                'links': parser.links,

                'images': parser.images

            }

        except Exception as e:

            logger.debug(f"HTML metadata extraction failed: {e}")

            return {}

    

    def get_supported_formats(self) -> Dict[str, Any]:

        """Return supported formats and compatibility information"""

        return {

            'success': True,

            'supported_formats': list(self.FORMAT_MIME_TYPES.keys()),

            'compatibility_matrix': self.FORMAT_COMPATIBILITY,

            'max_file_size_mb': int(config.MAX_DOC_SIZE / (1024 * 1024)),

            'converter': 'pypandoc (Pandoc)',

            'system_requirement': 'Pandoc system package',

            'pandoc_available': self.pandoc_available

        }

    

    def get_pandoc_info(self) -> Dict[str, Any]:

        """Get detailed information about installed Pandoc"""

        try:

            if not self.pandoc_available:

                return {

                    'status': 'not_available',

                    'error': 'Pandoc is not installed or not accessible'

                }

            

            version = pypandoc.get_pandoc_version()

            

            # Get list of supported input formats

            try:

                input_result = subprocess.run(

                    ['pandoc', '--list-input-formats'],

                    capture_output=True,

                    text=True,

                    timeout=5

                )

                input_formats = input_result.stdout.strip().split() if input_result.returncode == 0 else []

            except Exception as e:

                logger.warning(f"Could not get input formats: {e}")

                input_formats = []

            

            # Get list of supported output formats

            try:

                output_result = subprocess.run(

                    ['pandoc', '--list-output-formats'],

                    capture_output=True,

                    text=True,

                    timeout=5

                )

                output_formats = output_result.stdout.strip().split() if output_result.returncode == 0 else []

            except Exception as e:

                logger.warning(f"Could not get output formats: {e}")

                output_formats = []

            

            return {

                'status': 'installed',

                'version': str(version),

                'supported_input_formats': input_formats,

                'supported_output_formats': output_formats,

                'total_input_formats': len(input_formats),

                'total_output_formats': len(output_formats)

            }

        except Exception as e:

            logger.error(f"Failed to get Pandoc info: {e}")

            return {

                'status': 'error',

                'error': str(e)

            }
# --- END of DocumentProcessor ---


# --- Image Processing Service ---
class ImageProcessor:
    def __init__(self):
        self.output_dir = config.OUTPUT_DIR / "images"
        ensure_directory(self.output_dir)
    
    async def process_image(self, image_path: str, operations: List[str], **kwargs) -> Dict[str, Any]:
        """Process image with specified operations"""
        try:
            source = Path(image_path)
            if not source.exists():
                source = config.BASE_DIR / source.name
                if not source.exists():
                    raise FileNotFoundError(f"Source file not found: {image_path}")
            
            validate_file_size(source, config.MAX_IMG_SIZE)
            
            output_format = kwargs.get('output_format', 'png')
            quality = kwargs.get('quality', 85)
            width = kwargs.get('width')
            height = kwargs.get('height')
            
            # Generate output path
            output_name = f"{source.stem}_processed.{output_format}"
            output_path = self.output_dir / output_name
            
            # Process image in an executor to be non-blocking
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(
                None, 
                self._run_image_processing, 
                source, output_path, operations, 
                width, height, output_format, quality
            )
            
            metadata = await self._get_image_metadata(output_path)
            
            result = {
                "success": True,
                "message": f"Image processed with {len(operations)} operations",
                "output_path": str(output_path),
                "operations_applied": operations,
                "output_format": output_format,
                "metadata": metadata
            }
            log_to_history("process_image", "success", str(source), str(output_path), metadata=result)
            return result
            
        except Exception as e:
            logger.exception(f"Image processing failed")
            log_to_history("process_image", "error", str(image_path), error_message=str(e))
            return {
                "success": False,
                "error": str(e),
                "image_path": image_path,
                "operations": operations
            }
    
    def _run_image_processing(self, source: Path, output_path: Path, operations: List[str], width: Optional[int], height: Optional[int], output_format: str, quality: int):
        """Synchronous helper for image processing."""
        with Image.open(source) as img:
            processed_img = img.copy()
            
            # Apply operations
            for operation in operations:
                if operation == "resize" and (width or height):
                    processed_img = self._resize_image(processed_img, width, height)
                elif operation == "enhance":
                    processed_img = self._enhance_image(processed_img)
                elif operation == "convert":
                    pass  # Handled in save
            
            # Convert mode if necessary
            if processed_img.mode in ('RGBA', 'LA', 'P') and output_format.lower() == 'jpeg':
                background = Image.new('RGB', processed_img.size, (255, 255, 255))
                if processed_img.mode in ('RGBA', 'LA'):
                    background.paste(processed_img, mask=processed_img.split()[-1])
                else:
                    background.paste(processed_img)
                processed_img = background
            
            # Save processed image
            save_kwargs = {'format': output_format.upper()}
            if output_format.lower() in ['jpeg', 'webp']:
                save_kwargs.update({'quality': quality, 'optimize': True})
            
            processed_img.save(output_path, **save_kwargs)

    def _resize_image(self, img: Image.Image, width: Optional[int], height: Optional[int]) -> Image.Image:
        """Resize image maintaining aspect ratio"""
        if not width and not height:
            return img
        
        original_width, original_height = img.size
        
        if width and height:
            new_size = (width, height)
        elif width:
            ratio = width / original_width
            new_size = (width, int(original_height * ratio))
        else:  # height only
            ratio = height / original_height
            new_size = (int(original_width * ratio), height)
        
        return img.resize(new_size, Image.Resampling.LANCZOS)
    
    def _enhance_image(self, img: Image.Image) -> Image.Image:
        """Apply basic image enhancements"""
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.1)
        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(1.1)
        return img
    
    async def _get_image_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract image metadata"""
        try:
            # Run PIL in executor to be safe
            loop = asyncio.get_event_loop()
            stats = filepath.stat()
            
            def get_pil_meta():
                with Image.open(filepath) as img:
                    return {
                        "width": img.width,
                        "height": img.height,
                        "format": img.format,
                        "mode": img.mode,
                        "has_exif": hasattr(img, '_getexif') and img._getexif() is not None
                    }
            
            pil_meta = await loop.run_in_executor(None, get_pil_meta)
            pil_meta["file_size"] = stats.st_size
            return pil_meta
            
        except Exception:
            stats = filepath.stat()
            return {"file_size": stats.st_size, "format": "unknown"}

# --- Email Processing Service (UPGRADED) ---
class EmailProcessor:
    def __init__(self):
        self.output_dir = config.OUTPUT_DIR / "email"
        ensure_directory(self.output_dir)
        
        # --- NEW: Configure Gemini API for summarization ---
        self.genai_model = None
        try:
            api_key = os.getenv('GEMINI_API_KEY')
            if not api_key:
                logger.warning("GEMINI_API_KEY not set in .env. Email summarization will fail.")
            else:
                genai.configure(api_key=api_key)
                self.genai_model = genai.GenerativeModel(os.getenv('GEMINI_MODEL', 'gemini-1.5-flash'))
                logger.info("EmailProcessor configured for Gemini summarization.")
        except Exception as e:
            logger.error(f"Failed to configure Gemini for EmailProcessor: {e}")

    async def _summarize_text_with_gemini(self, text_content: str) -> str:
        """Helper to call Gemini API for summarization."""
        if not self.genai_model:
            raise RuntimeError("Gemini API is not configured on the server for summarization.")
        
        prompt = f"""
        You are a helpful assistant. Please provide a concise, bullet-pointed summary
        of the following email content. Focus on key topics, questions, and action items.
        
        EMAIL CONTENT:
        ---
        {text_content}
        ---
        
        SUMMARY:
        """
        
        try:
            # Use run_in_executor for the blocking genai call
            response = await asyncio.get_event_loop().run_in_executor(
                None, self.genai_model.generate_content, prompt
            )
            return response.text.strip()
        except Exception as e:
            logger.error(f"Gemini summarization failed: {e}")
            raise RuntimeError(f"AI summarization failed: {e}")

    # --- NEW TOOL ---
    async def summarize_emails(self, archive_path: str, **kwargs) -> Dict[str, Any]:
        """Summarizes the content of an email archive."""
        try:
            archive = Path(archive_path)
            if not archive.exists():
                archive = config.BASE_DIR / archive.name
                if not archive.exists():
                    raise FileNotFoundError(f"Email archive not found: {archive_path}")

            validate_file_size(archive, config.MAX_EMAIL_SIZE)
            
            loop = asyncio.get_event_loop()
            emails = await loop.run_in_executor(None, self._parse_email_archive_sync, archive)
            
            if not emails:
                raise ValueError("No emails found in archive to summarize.")

            # Combine all email content into one large text block
            full_text_content = ""
            for email_data in emails:
                full_text_content += f"--- EMAIL FROM: {email_data.get('sender', 'N/A')} | SUBJECT: {email_data.get('subject', 'N/A')} ---\n"
                full_text_content += email_data.get('content', '') + "\n\n"

            # Summarize the combined text
            summary = await self._summarize_text_with_gemini(full_text_content)
            
            # --- NEW: Save summary to .txt file ---
            output_name = f"{archive.stem}_summary.txt"
            output_path = self.output_dir / output_name
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"Summary for: {archive_path}\n")
                f.write(f"Total Emails Found: {len(emails)}\n")
                f.write("="*30 + "\n\n")
                f.write(summary)
            
            result = {
                "success": True,
                "message": f"Successfully summarized {len(emails)} email(s). Saved to {output_path.name}",
                "emails_found": len(emails),
                "summary": summary, # Also return summary in response
                "output_path": str(output_path), # Return the new file path
                "source_path": str(archive_path)
            }
            log_to_history("summarize_emails", "success", str(archive), str(output_path), metadata={"emails_found": len(emails)})
            return result
            
        except Exception as e:
            logger.exception(f"Email summarization failed")
            log_to_history("summarize_emails", "error", str(archive_path), error_message=str(e))
            return {
                "success": False,
                "error": str(e),
                "archive_path": archive_path
            }
    
    async def search_emails(self, archive_path: str, query: str, **kwargs) -> Dict[str, Any]:
        """Search emails in archive"""
        try:
            archive = Path(archive_path)
            if not archive.exists():
                archive = config.BASE_DIR / archive.name
                if not archive.exists():
                    raise FileNotFoundError(f"Email archive not found: {archive_path}")
            
            validate_file_size(archive, config.MAX_EMAIL_SIZE)
            
            loop = asyncio.get_event_loop()
            emails = await loop.run_in_executor(None, self._parse_email_archive_sync, archive)
            
            query_lower = query.lower()
            matching_emails = []
            
            for email_data in emails:
                if self._email_matches_query(email_data, query_lower):
                    matching_emails.append({
                        "subject": email_data.get("subject", ""),
                        "sender": email_data.get("sender", ""),
                        "date": email_data.get("date", ""),
                        "snippet": email_data.get("content", "")[:200] + "..."
                    })
            
            result = {
                "success": True,
                "message": f"Found {len(matching_emails)} matching emails",
                "total_emails": len(emails),
                "matching_emails": len(matching_emails),
                "results": matching_emails[:50],  # Limit results
                "query": query
            }
            log_to_history("search_emails", "success", str(archive), metadata={"query": query, "matches": len(matching_emails)})
            return result
            
        except Exception as e:
            logger.exception(f"Email search failed")
            log_to_history("search_emails", "error", str(archive_path), error_message=str(e), metadata={"query": query})
            return {
                "success": False,
                "error": str(e),
                "archive_path": archive_path,
                "query": query
            }
    
    async def analyze_communication_patterns(self, archive_path: str, **kwargs) -> Dict[str, Any]:
        """Analyze communication patterns"""
        try:
            archive = Path(archive_path)
            if not archive.exists():
                archive = config.BASE_DIR / archive.name
                if not archive.exists():
                    raise FileNotFoundError(f"Email archive not found: {archive_path}")

            loop = asyncio.get_event_loop()
            emails = await loop.run_in_executor(None, self._parse_email_archive_sync, archive)
            
            if not emails:
                raise ValueError("No emails found in archive")
            
            sent_count = sum(1 for e in emails if self._is_sent_email(e))
            received_count = len(emails) - sent_count
            
            contacts = {}
            for email_data in emails:
                sender = email_data.get("sender", "")
                if sender and sender not in contacts:
                    contacts[sender] = {"count": 0, "first_contact": None, "last_contact": None}
                
                contacts[sender]["count"] += 1
                date_str = email_data.get("date", "")
                contacts[sender]["last_contact"] = date_str
                if not contacts[sender]["first_contact"]:
                    contacts[sender]["first_contact"] = date_str
            
            top_contacts = sorted(contacts.items(), key=lambda x: x[1]["count"], reverse=True)[:10]
            
            result = {
                "success": True,
                "message": f"Analyzed {len(emails)} emails",
                "total_emails": len(emails),
                "sent_count": sent_count,
                "received_count": received_count,
                "unique_contacts": len(contacts),
                "top_contacts": [{"email": email, **data} for email, data in top_contacts],
                "analysis_date": datetime.now().isoformat()
            }
            log_to_history("analyze_communication_patterns", "success", str(archive), metadata={"total_emails": len(emails)})
            return result
            
        except Exception as e:
            logger.exception(f"Communication analysis failed")
            log_to_history("analyze_communication_patterns", "error", str(archive_path), error_message=str(e))
            return {
                "success": False,
                "error": str(e),
                "archive_path": archive_path
            }
    
    def _parse_email_archive_sync(self, archive_path: Path) -> List[Dict[str, Any]]:
        """Synchronous helper to parse email archive file"""
        emails = []
        try:
            if archive_path.suffix.lower() == '.mbox':
                mbox = mailbox.mbox(str(archive_path))
                for message in mbox:
                    email_data = self._parse_email_message(message)
                    if email_data:
                        emails.append(email_data)
            elif archive_path.suffix.lower() == '.eml':
                with open(archive_path, 'rb') as f:
                    message = BytesParser(policy=default).parse(f)
                    email_data = self._parse_email_message(message)
                    if email_data:
                        emails.append(email_data)
        except Exception as e:
            logger.warning(f"Failed to parse email archive: {e}")
        return emails
    
    def _parse_email_message(self, message) -> Optional[Dict[str, Any]]:
        """Parse individual email message"""
        try:
            return {
                "subject": message.get("Subject", ""),
                "sender": message.get("From", ""),
                "recipients": message.get("To", ""),
                "date": message.get("Date", ""),
                "content": self._get_email_content(message),
                "message_id": message.get("Message-ID", "")
            }
        except Exception:
            return None
    
    def _get_email_content(self, message) -> str:
        """Extract email content"""
        try:
            if message.is_multipart():
                for part in message.walk():
                    if part.get_content_type() == "text/plain":
                        return part.get_content()
            else:
                return message.get_content()
        except Exception:
            return ""
    
    def _email_matches_query(self, email_data: Dict[str, Any], query: str) -> bool:
        """Check if email matches search query"""
        searchable_text = " ".join([
            email_data.get("subject", ""),
            email_data.get("sender", ""),
            email_data.get("content", "")
        ]).lower()
        return query in searchable_text
    
    def _is_sent_email(self, email_data: Dict[str, Any]) -> bool:
        """Simple heuristic to determine if email was sent"""
        sender = email_data.get("sender", "").lower()
        return "noreply" not in sender and "no-reply" not in sender

# --- News Processing Service (UPGRADED) ---
class NewsProcessor:
    def __init__(self):
        self.output_dir = config.OUTPUT_DIR / "news"
        ensure_directory(self.output_dir)
    
    async def add_news_source(self, feed_url: str, category: str = "general", **kwargs) -> Dict[str, Any]:
        """Add RSS news source to the database"""
        try:
            # Run blocking feedparser in executor
            loop = asyncio.get_event_loop()
            feed = await loop.run_in_executor(None, feedparser.parse, feed_url)

            if feed.bozo:
                raise ValueError(f"Invalid RSS feed: {feed_url}")
            
            source_data = {
                "url": feed_url,
                "title": feed.feed.get("title", "Unknown"),
                "description": feed.feed.get("description", ""),
                "category": category,
            }
            
            with get_db_connection() as conn:
                # --- START: SQL BUG FIX ---
                # This query now matches the 8-column schema from init_database()
                conn.execute(
                    """
                    INSERT INTO news_sources (url, title, description, category, last_updated)
                    VALUES (?, ?, ?, ?, ?)
                    ON CONFLICT(url) DO UPDATE SET
                        title=excluded.title,
                        description=excluded.description,
                        category=excluded.category,
                        last_updated=?
                    """,
                    (source_data["url"], source_data["title"], source_data["description"], source_data["category"], datetime.now().isoformat())
                )
                # --- END: SQL BUG FIX ---
            
            log_to_history("add_news_source", "success", feed_url, metadata=source_data)
            return {
                "success": True,
                "message": f"Added/Updated news source: {source_data['title']}",
                "source_data": source_data
            }
            
        except Exception as e:
            logger.exception(f"Failed to add news source")
            log_to_history("add_news_source", "error", feed_url, error_message=str(e))
            return {
                "success": False,
                "error": str(e),
                "feed_url": feed_url
            }
    
    async def get_trending_topics(self, time_window: str = "24h", **kwargs) -> Dict[str, Any]:
        """Get trending topics from database sources"""
        all_articles = []
        source_count = 0
        
        # --- NEW: Get optional keyword from arguments ---
        keyword_filter = kwargs.get("category_keyword", "").lower()
        loop = asyncio.get_event_loop()

        try:
            with get_db_connection() as conn:
                sources = conn.execute("SELECT id, url, title FROM news_sources WHERE status='active'").fetchall()
                source_count = len(sources)
            
                for source in sources:
                    try:
                        # Run blocking feedparser in executor
                        feed = await loop.run_in_executor(None, feedparser.parse, source["url"])
                        
                        for entry in feed.entries[:20]:
                            article = {
                                "title": entry.get("title", ""),
                                "summary": entry.get("summary", ""),
                                "published": entry.get("published", ""),
                                "source": source["title"],
                                "link": entry.get("link", ""),
                                "source_id": source["id"],
                                "hash": get_file_hash(entry.get("link", "") + entry.get("title", ""))
                            }
                            all_articles.append(article)
                            
                            conn.execute(
                                """
                                INSERT INTO news_articles (source_id, title, summary, url, published_date, hash)
                                VALUES (?, ?, ?, ?, ?, ?)
                                ON CONFLICT(hash) DO NOTHING
                                """,
                                (article["source_id"], article["title"], article["summary"], article["link"], article["published"], article["hash"])
                            )
                    except Exception as e:
                        logger.warning(f"Failed to fetch from {source['url']}: {e}")
            
            word_counts = {}
            # --- NEW: Filter articles if keyword is provided ---
            if keyword_filter:
                filtered_articles = [
                    article for article in all_articles 
                    if keyword_filter in article["title"].lower() or keyword_filter in article["summary"].lower()
                ]
            else:
                filtered_articles = all_articles

            for article in filtered_articles:
                words = re.findall(r'\b\w+\b', article["title"].lower())
                for word in words:
                    if len(word) > 4 and word not in ['this', 'that', 'with', 'from', 'news', keyword_filter]:
                        word_counts[word] = word_counts.get(word, 0) + 1
            
            trending_words = sorted(word_counts.items(), key=lambda x: x[1], reverse=True)[:15]
            
            with get_db_connection() as conn:
                for word, count in trending_words:
                    conn.execute(
                        """
                        INSERT INTO trending_keywords (keyword, count, category, time_window)
                        VALUES (?, ?, ?, ?)
                        ON CONFLICT(keyword, category, time_window) DO UPDATE SET
                            count=excluded.count,
                            last_seen=CURRENT_TIMESTAMP
                        """,
                        (word, count, keyword_filter if keyword_filter else 'general', time_window)
                    )
            
            # --- NEW: Create .txt output as requested ---
            output_txt_path = self.output_dir / f"{keyword_filter if keyword_filter else 'latest'}_news.txt"
            with open(output_txt_path, 'w', encoding='utf-8') as f:
                f.write(f"News Report for: '{keyword_filter if keyword_filter else 'All Topics'}'\n")
                f.write(f"Found {len(filtered_articles)} matching articles.\n")
                f.write("="*30 + "\n\n")
                for article in filtered_articles[:30]: # Write top 30 articles to txt
                    f.write(f"Title: {article['title']}\n")
                    f.write(f"Source: {article['source']}\n")
                    f.write(f"Link: {article['link']}\n")
                    f.write(f"Published: {article['published']}\n")
                    f.write(f"Summary: {article['summary']}\n\n")
                    f.write("-"*30 + "\n\n")

            result = {
                "success": True,
                "message": f"Found {len(filtered_articles)} articles. Report saved to {output_txt_path.name}",
                "total_articles_scanned": len(all_articles),
                "matching_articles": len(filtered_articles),
                "trending_keywords": [{"word": word, "mentions": count} for word, count in trending_words],
                "recent_articles": filtered_articles[:10], # Show 10 in JSON
                "time_window": time_window,
                "output_path": str(output_txt_path) # Add the new output path
            }
            
            # Save the JSON output as well
            output_json_path = self.output_dir / "latest_trending.json"
            with open(output_json_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2, default=str)

            log_to_history("get_trending_topics", "success", metadata={"articles_found": len(all_articles), "sources": source_count})
            return result
            
        except Exception as e:
            logger.exception(f"Failed to get trending topics")
            log_to_history("get_trending_topics", "error", error_message=str(e))
            return {
                "success": False,
                "error": str(e),
                "time_window": time_window
            }

# Initialize processors
doc_processor = DocumentProcessor()
img_processor = ImageProcessor()
email_processor = EmailProcessor()
news_processor = NewsProcessor()

# --- MCP Tool Definitions ---
@app.list_tools()
async def list_tools() -> List[types.Tool]:
    """List all available tools"""
    return [
        types.Tool(
            name="convert_document",
            description="Convert document between formats (PDF, DOCX, PPTX, HTML, TXT)",
            inputSchema={
                "type": "object",
                "properties": {
                    "source_path": {"type": "string", "description": "Path to source document"},
                    "target_format": {"type": "string", "enum": config.DOC_FORMATS, "description": "Target format"},
                },
                "required": ["source_path", "target_format"]
            }
        ),
        types.Tool(
            name="process_image",
            description="Process image with resize, enhance, and format conversion",
            inputSchema={
                "type": "object",
                "properties": {
                    "image_path": {"type": "string", "description": "Path to source image"},
                    "operations": {"type": "array", "items": {"type": "string", "enum": ["resize", "enhance", "convert"]}, "description": "Operations to apply"},
                    "output_format": {"type": "string", "enum": config.IMG_FORMATS, "default": "png"},
                    "quality": {"type": "integer", "minimum": 1, "maximum": 100, "default": 85},
                    "width": {"type": "integer", "minimum": 1},
                    "height": {"type": "integer", "minimum": 1}
                },
                "required": ["image_path", "operations"]
            }
        ),
        types.Tool(
            name="search_emails",
            description="Search through email archives (mbox/eml)",
            inputSchema={
                "type": "object",
                "properties": {
                    "archive_path": {"type": "string", "description": "Path to email archive (mbox/eml)"},
                    "query": {"type": "string", "description": "Search query"},
                },
                "required": ["archive_path", "query"]
            }
        ),
        types.Tool(
            name="analyze_communication_patterns",
            description="Analyze email communication patterns from an archive",
            inputSchema={
                "type": "object",
                "properties": {
                    "archive_path": {"type": "string", "description": "Path to email archive"},
                },
                "required": ["archive_path"]
            }
        ),
        
        # --- NEW TOOL DEFINITION ---
        types.Tool(
            name="summarize_emails",
            description="Reads an entire email archive (.mbox or .eml) and generates a concise summary of all content.",
            inputSchema={
                "type": "object",
                "properties": {
                    "archive_path": {"type": "string", "description": "Path to email archive (mbox/eml)"},
                },
                "required": ["archive_path"]
            }
        ),
        # --- END NEW TOOL ---

        types.Tool(
            name="add_news_source",
            description="Add RSS/news source for monitoring",
            inputSchema={
                "type": "object",
                "properties": {
                    "feed_url": {"type": "string", "description": "RSS feed URL"},
                    "category": {"type": "string", "default": "general"},
                },
                "required": ["feed_url"]
            }
        ),
        types.Tool(
            name="get_trending_topics",
            description="Get trending topics from news sources. Can optionally filter by a category keyword.",
            inputSchema={
                "type": "object",
                "properties": {
                    "time_window": {"type": "string", "enum": ["1h", "6h", "24h", "7d"], "default": "24h"},
                    "category_keyword": {"type": "string", "description": "A keyword to filter articles by (e.g., 'sports')"}
                }
            }
        ),
        types.Tool(
            name="list_supported_formats",
            description="Get list of supported formats for each service",
            inputSchema={
                "type": "object",
                "properties": {
                    "service": {"type": "string", "enum": ["document", "image", "email", "all"], "default": "all"}
                }
            }
        )
    ]

@app.call_tool()
async def call_tool(name: str, arguments: dict) -> List[types.TextContent]:
    """Handle tool calls"""
    try:
        result = None
        
        if name == "convert_document":
            result = await doc_processor.convert_document(**arguments)
        elif name == "process_image":
            result = await img_processor.process_image(**arguments)
        elif name == "search_emails":
            result = await email_processor.search_emails(**arguments)
        elif name == "analyze_communication_patterns":
            result = await email_processor.analyze_communication_patterns(**arguments)
        
        # --- NEW TOOL HANDLER ---
        elif name == "summarize_emails":
            result = await email_processor.summarize_emails(**arguments)
        # --- END NEW TOOL ---

        elif name == "add_news_source":
            result = await news_processor.add_news_source(**arguments)
        elif name == "get_trending_topics":
            result = await news_processor.get_trending_topics(**arguments)
        elif name == "list_supported_formats":
            service = arguments.get("service", "all")
            formats = {}
            if service in ["document", "all"]:
                formats["document"] = config.DOC_FORMATS
            if service in ["image", "all"]:
                formats["image"] = config.IMG_FORMATS
            if service in ["email", "all"]:
                formats["email"] = config.EMAIL_FORMATS
            result = { "success": True, "supported_formats": formats, "service": service }
        else:
            result = { "success": False, "error": f"Unknown tool: {name}" }
        
        return [types.TextContent(
            type="text",
            text=json.dumps(result, indent=2, default=str)
        )]
            
    except Exception as e:
        logger.exception(f"Tool call failed: {name}")
        return [types.TextContent(
            type="text",
            text=json.dumps({
                "success": False,
                "error": str(e),
                "tool": name,
                "arguments": arguments
            }, indent=2, default=str)
        )]

# --- Resources ---
@app.list_resources()
async def list_resources() -> List[types.Resource]:
    """List available resources"""
    resources = []
    
    # List output files as resources
    for service_dir in ["documents", "images", "email", "news"]:
        service_path = config.OUTPUT_DIR / service_dir
        if service_path.exists():
            for file_path in service_path.iterdir():
                if file_path.is_file():
                    resources.append(types.Resource(
                        uri=f"file://{file_path.resolve()}", # Use resolve() for a full path
                        name=f"{service_dir}/{file_path.name}",
                        description=f"Processed {service_dir[:-1]} file",
                        mimeType=_get_mime_type(file_path)
                    ))
    
    return resources

@app.read_resource()
async def read_resource(uri: str) -> str:
    """Read resource content"""
    try:
        if uri.startswith("file://"):
            file_path = Path(uri[7:])  # Remove file:// prefix
            
            # Security check - ensure file is in output directory
            if not file_path.resolve().is_relative_to(config.OUTPUT_DIR.resolve()):
                 raise ValueError("Access denied: File outside output directory")
            
            if not file_path.exists():
                raise FileNotFoundError(f"Resource not found: {uri}")
            
            # Read file content
            if _is_text_file(file_path):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                # For binary files, return base64 encoded content
                import base64
                with open(file_path, 'rb') as f:
                    content = base64.b64encode(f.read()).decode('utf-8')
                return f"data:{_get_mime_type(file_path)};base64,{content}"
        
        else:
            raise ValueError(f"Unsupported URI scheme: {uri}")
            
    except Exception as e:
        logger.error(f"Failed to read resource {uri}: {e}")
        raise

def _get_mime_type(file_path: Path) -> str:
    """Get MIME type for file"""
    try:
        return magic.from_file(str(file_path), mime=True)
    except:
        # Fallback based on extension
        ext = file_path.suffix.lower()
        mime_map = {
            '.txt': 'text/plain',
            '.html': 'text/html',
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.webp': 'image/webp',
            '.json': 'application/json'
        }
        return mime_map.get(ext, 'application/octet-stream')

def _is_text_file(file_path: Path) -> bool:
    """Check if file is text-based"""
    text_extensions = {'.txt', '.html', '.htm', '.md', '.json', '.xml', '.csv'}
    return file_path.suffix.lower() in text_extensions

# --- Prompts ---
@app.list_prompts()
async def list_prompts() -> List[types.Prompt]:
    """List available prompts"""
    return [
        types.Prompt(
            name="system_status",
            description="Get system status and statistics",
            arguments=[]
        ),
        types.Prompt(
            name="processing_summary",
            description="Get summary of recent processing activities",
            arguments=[
                types.PromptArgument(
                    name="time_period",
                    description="Time period for summary (1h, 24h, 7d)",
                    required=False
                )
            ]
        )
    ]

@app.get_prompt()
async def get_prompt(name: str, arguments: dict) -> types.GetPromptResult:
    """Get prompt content"""
    try:
        if name == "system_status":
            status_info = await _get_system_status()
            return types.GetPromptResult(
                description="Current system status",
                messages=[
                    types.PromptMessage(
                        role="user",
                        content=types.TextContent(
                            type="text",
                            text=f"System Status Report:\n\n{json.dumps(status_info, indent=2, default=str)}"
                        )
                    )
                ]
            )
        
        elif name == "processing_summary":
            time_period = arguments.get("time_period", "24h")
            summary_info = await _get_processing_summary(time_period)
            return types.GetPromptResult(
                description=f"Processing summary for {time_period}",
                messages=[
                    types.PromptMessage(
                        role="user",
                        content=types.TextContent(
                            type="text",
                            text=f"Processing Summary ({time_period}):\n\n{json.dumps(summary_info, indent=2, default=str)}"
                        )
                    )
                ]
            )
        
        else:
            raise ValueError(f"Unknown prompt: {name}")
    
    except Exception as e:
        logger.error(f"Failed to get prompt {name}: {e}")
        return types.GetPromptResult(
            description="Error getting prompt",
            messages=[
                types.PromptMessage(
                    role="user",
                    content=types.TextContent(
                        type="text",
                        text=f"Error: {str(e)}"
                    )
                )
            ]
        )

async def _get_system_status() -> Dict[str, Any]:
    """Get current system status"""
    try:
        # Count files in output directories
        file_counts = {}
        total_size = 0
        
        for service_dir in ["documents", "images", "email", "news"]:
            service_path = config.OUTPUT_DIR / service_dir
            if service_path.exists():
                files = list(service_path.iterdir())
                file_counts[service_dir] = len(files)
                total_size += sum(f.stat().st_size for f in files if f.is_file())
            else:
                file_counts[service_dir] = 0
        
        # Check disk space
        disk_usage = shutil.disk_usage(config.BASE_DIR)
        
        # Get DB stats
        db_stats = {}
        try:
            with get_db_connection() as conn:
                db_stats["news_sources"] = conn.execute("SELECT COUNT(*) FROM news_sources").fetchone()[0]
                db_stats["news_articles"] = conn.execute("SELECT COUNT(*) FROM news_articles").fetchone()[0]
                db_stats["trending_keywords"] = conn.execute("SELECT COUNT(*) FROM trending_keywords").fetchone()[0]
                db_stats["processing_history"] = conn.execute("SELECT COUNT(*) FROM processing_history").fetchone()[0]
        except Exception as db_e:
            logger.error(f"Database stats failed: {db_e}")
            db_stats = {"error": str(db_e)}
            
        return {
            "status": "running",
            "timestamp": datetime.now().isoformat(),
            "directories": {
                "base": str(config.BASE_DIR),
                "temp": str(config.TEMP_DIR),
                "output": str(config.OUTPUT_DIR),
                "cache": str(config.CACHE_DIR)
            },
            "file_counts": file_counts,
            "total_output_size": total_size,
            "disk_usage": {
                "total": disk_usage.total,
                "used": disk_usage.used,
                "free": disk_usage.free,
                "percent_used": (disk_usage.used / disk_usage.total) * 100
            },
            "database_stats": db_stats,
            "supported_formats": {
                "documents": config.DOC_FORMATS,
                "images": config.IMG_FORMATS,
                "email": config.EMAIL_FORMATS
            }
        }
        
    except Exception as e:
        return {
            "status": "error",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

async def _get_processing_summary(time_period: str) -> Dict[str, Any]:
    """Get processing summary for time period from DB"""
    try:
        cutoff_hours = {
            "1h": 1,
            "24h": 24,
            "7d": 168
        }.get(time_period, 24)
        
        cutoff_time = datetime.now().timestamp() - (cutoff_hours * 3600)
        cutoff_iso = datetime.fromtimestamp(cutoff_time).isoformat()

        summary = {}
        with get_db_connection() as conn:
            rows = conn.execute(
                """
                SELECT tool_name, status, COUNT(*) as count
                FROM processing_history
                WHERE created_at >= ?
                GROUP BY tool_name, status
                """,
                (cutoff_iso,)
            ).fetchall()

            for row in rows:
                if row["tool_name"] not in summary:
                    summary[row["tool_name"]] = {"success": 0, "error": 0}
                summary[row["tool_name"]][row["status"]] = row["count"]
        
        total_ops = sum(s["success"] + s["error"] for s in summary.values())
        
        return {
            "time_period": time_period,
            "summary_generated": datetime.now().isoformat(),
            "processing_summary": summary,
            "total_operations": total_ops,
        }
        
    except Exception as e:
        return {
            "error": str(e),
            "time_period": time_period,
            "timestamp": datetime.now().isoformat()
        }

# Cleanup function
async def cleanup_temp_files():
    """Clean up temporary files"""
    try:
        if config.TEMP_DIR.exists():
            for file_path in config.TEMP_DIR.iterdir():
                if file_path.is_file():
                    # Remove files older than 1 hour
                    if file_path.stat().st_mtime < datetime.now().timestamp() - 3600:
                        file_path.unlink()
                        logger.info(f"Cleaned up temp file: {file_path}")
    except Exception as e:
        logger.error(f"Cleanup failed: {e}")

# Main server function
async def main():
    """Main server function"""
    logger.info("Starting Content & Media Processing MCP Server...")
    
    # Perform initial setup
    try:
        # --- START: BUG FIX ---
        # Initialize database *before* initializing processors
        init_database()
        # --- END: BUG FIX ---
        
        # Create required directories
        config.__post_init__()
        
        # Initialize processors
        global doc_processor, img_processor, email_processor, news_processor
        doc_processor = DocumentProcessor()
        img_processor = ImageProcessor()
        email_processor = EmailProcessor()
        news_processor = NewsProcessor()
        
        # Schedule cleanup task
        asyncio.create_task(_periodic_cleanup())
        
        logger.info("Server initialization complete")
        
        # Run the stdio server
        async with stdio_server() as (read_stream, write_stream):
            await app.run(
                read_stream,
                write_stream,
                app.create_initialization_options()
            )
            
    except Exception as e:
        logger.error(f"Server startup failed: {e}")
        sys.exit(1)

async def _periodic_cleanup():
    """Periodic cleanup task"""
    while True:
        try:
            await asyncio.sleep(3600)  # Run every hour
            await cleanup_temp_files()
        except Exception as e:
            logger.error(f"Periodic cleanup failed: {e}")

# Entry point
if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        logger.info("Server shutdown requested")
    except Exception as e:
        logger.error(f"Server error: {e}")
        sys.exit(1)